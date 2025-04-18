import streamlit as st
import pandas as pd
import tempfile
import os
from pathlib import Path
import requests
import json
import base64
from gtts import gTTS
import PyPDF2
import docx
from io import BytesIO
import time
import re
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
import datetime
import uuid
import logging
import hashlib
import csv
import plotly.express as px
import plotly.graph_objects as go
import zipfile

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("document_analyzer.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("DocumentAnalyzer")

# Page configuration
st.set_page_config(
    page_title="Advanced Document Analyzer",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        margin-bottom: 1rem;
    }
    .subheader {
        font-size: 1.5rem;
        margin-bottom: 1rem;
    }
    .summary-container {
        background-color: #f0f2f6;
        padding: 20px;
        border-radius: 10px;
        margin-top: 20px;
    }
    .stAudio {
        margin-top: 20px;
    }
    .search-result {
        background-color: #e6f3ff;
        padding: 15px;
        border-radius: 5px;
        margin-bottom: 10px;
        border-left: 4px solid #0066cc;
    }
    .highlight {
        background-color: #ffff00;
        padding: 2px;
        border-radius: 3px;
    }
    .log-entry {
        padding: 8px;
        border-bottom: 1px solid #eee;
        font-size: 0.9em;
    }
    .log-info {
        color: #0066cc;
    }
    .log-warning {
        color: #ff9900;
    }
    .log-error {
        color: #cc0000;
    }
    .metrics-card {
        background-color: white;
        padding: 15px;
        border-radius: 5px;
        box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        text-align: center;
    }
    .metrics-value {
        font-size: 1.8rem;
        font-weight: bold;
        margin: 10px 0;
    }
    .metrics-label {
        font-size: 0.9rem;
        color: #666;
    }
    .history-item {
        padding: 10px;
        border: 1px solid #ddd;
        border-radius: 5px;
        margin-bottom: 8px;
        cursor: pointer;
    }
    .history-item:hover {
        background-color: #f5f5f5;
    }
    .tab-content {
        padding: 20px 0;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state variables
def init_session_state():
    if 'extracted_text' not in st.session_state:
        st.session_state.extracted_text = ""
    if 'paragraphs' not in st.session_state:
        st.session_state.paragraphs = []
    if 'search_results' not in st.session_state:
        st.session_state.search_results = []
    if 'document_history' not in st.session_state:
        st.session_state.document_history = []
    if 'activity_log' not in st.session_state:
        st.session_state.activity_log = []
    if 'current_file_id' not in st.session_state:
        st.session_state.current_file_id = None
    if 'analysis_results' not in st.session_state:
        st.session_state.analysis_results = {}
    if 'settings' not in st.session_state:
        st.session_state.settings = {
            'default_model': 'llama3',
            'summary_length': 'medium',
            'search_results_count': 5,
            'highlight_enabled': True,
            'save_logs': True,
            'dark_mode': False
        }

init_session_state()

# Log activity function
def log_activity(action, details="", level="INFO"):
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_entry = {
        "timestamp": timestamp,
        "action": action,
        "details": details,
        "level": level
    }
    
    # Add to session state
    st.session_state.activity_log.append(log_entry)
    
    # Log to file system
    if level == "INFO":
        logger.info(f"{action}: {details}")
    elif level == "WARNING":
        logger.warning(f"{action}: {details}")
    elif level == "ERROR":
        logger.error(f"{action}: {details}")
    
    return log_entry

# Function to calculate document hash
def calculate_file_hash(file_bytes):
    return hashlib.md5(file_bytes).hexdigest()

# Function to add document to history
def add_to_history(filename, file_size, file_hash, file_type, word_count):
    # Check if document already exists in history
    for item in st.session_state.document_history:
        if item['file_hash'] == file_hash:
            # Update access time and move to top
            item['last_accessed'] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            st.session_state.document_history.remove(item)
            st.session_state.document_history.insert(0, item)
            return item['file_id']
    
    # Create new history entry
    file_id = str(uuid.uuid4())
    history_item = {
        'file_id': file_id,
        'filename': filename,
        'file_size': file_size,
        'file_hash': file_hash,
        'file_type': file_type,
        'word_count': word_count,
        'upload_time': datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'last_accessed': datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }
    
    # Add to front of list
    st.session_state.document_history.insert(0, history_item)
    
    # Keep history to a reasonable size
    if len(st.session_state.document_history) > 20:
        st.session_state.document_history = st.session_state.document_history[:20]
    
    log_activity("Document Added", f"Added {filename} to history")
    return file_id

# Cache the text extraction functions to improve performance
@st.cache_data
def extract_text_from_pdf(file_bytes):
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        log_activity("PDF Extraction", f"Extracted {len(pdf_reader.pages)} pages")
        return text
    except Exception as e:
        error_msg = f"Error extracting text from PDF: {e}"
        log_activity("PDF Extraction Error", error_msg, "ERROR")
        st.error(error_msg)
        return ""

@st.cache_data
def extract_text_from_docx(file_bytes):
    try:
        doc = docx.Document(BytesIO(file_bytes))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        log_activity("DOCX Extraction", f"Extracted {len(doc.paragraphs)} paragraphs")
        return text
    except Exception as e:
        error_msg = f"Error extracting text from DOCX: {e}"
        log_activity("DOCX Extraction Error", error_msg, "ERROR")
        st.error(error_msg)
        return ""

@st.cache_data
def extract_text_from_txt(file_bytes):
    try:
        text = file_bytes.decode("utf-8")
        log_activity("TXT Extraction", f"Extracted {len(text.splitlines())} lines")
        return text
    except Exception as e:
        error_msg = f"Error extracting text from TXT: {e}"
        log_activity("TXT Extraction Error", error_msg, "ERROR")
        st.error(error_msg)
        return ""

# Function to split text into paragraphs or chunks
def split_into_paragraphs(text, min_length=50):
    # Split text by line breaks or paragraph markers
    raw_paragraphs = re.split(r'\n\s*\n|\r\n\s*\r\n', text)
    
    # Filter out very short paragraphs and clean them
    paragraphs = []
    for para in raw_paragraphs:
        cleaned = para.strip()
        if len(cleaned) >= min_length:  # Only keep paragraphs with meaningful content
            paragraphs.append(cleaned)
    
    # If paragraphs are still too few, split by sentences
    if len(paragraphs) < 5:
        all_sentences = re.split(r'(?<=[.!?])\s+', text)
        # Group sentences into chunks of 3-5
        chunk_size = 3
        paragraphs = [' '.join(all_sentences[i:i+chunk_size]) 
                      for i in range(0, len(all_sentences), chunk_size)
                      if ' '.join(all_sentences[i:i+chunk_size]).strip()]
    
    log_activity("Text Processing", f"Split document into {len(paragraphs)} paragraphs")
    return paragraphs

# Function to search for relevant content
def search_document(query, paragraphs, top_k=5):
    if not query or not paragraphs:
        return []
    
    try:
        # Create TF-IDF vectorizer
        vectorizer = TfidfVectorizer(stop_words='english')
        
        # Fit and transform paragraphs
        tfidf_matrix = vectorizer.fit_transform(paragraphs)
        
        # Transform the query
        query_vector = vectorizer.transform([query])
        
        # Calculate cosine similarity
        cosine_similarities = cosine_similarity(query_vector, tfidf_matrix).flatten()
        
        # Get top k results
        top_indices = cosine_similarities.argsort()[-top_k:][::-1]
        
        # Filter out results with very low similarity
        results = [
            {"text": paragraphs[i], "score": float(cosine_similarities[i]), "index": i}
            for i in top_indices
            if cosine_similarities[i] > 0.1  # Filtering threshold
        ]
        
        log_activity("Search", f"Query: '{query}' found {len(results)} results")
        return results
    except Exception as e:
        error_msg = f"Search error: {e}"
        log_activity("Search Error", error_msg, "ERROR")
        st.error(error_msg)
        return []

# Function to highlight search terms in text
def highlight_text(text, search_term):
    if not search_term or not st.session_state.settings['highlight_enabled']:
        return text
    
    # Escape special regex characters in search term
    escaped_term = re.escape(search_term)
    
    # Create a regex pattern that's case insensitive
    pattern = re.compile(f"({escaped_term})", re.IGNORECASE)
    
    # Replace with highlighted version
    highlighted = pattern.sub(r'<span class="highlight">\1</span>', text)
    return highlighted

# Function to get summary from Ollama
def get_summary_from_ollama(text, model="llama3", length="medium"):
    try:
        # Adjust summary length instructions
        length_instructions = {
            "short": "a brief summary (about 100 words)",
            "medium": "a concise summary (about 200 words)",
            "long": "a comprehensive summary (about 400 words)"
        }
        
        length_instr = length_instructions.get(length, length_instructions["medium"])
        
        # Prepare the prompt for summarization
        prompt = f"""Please provide {length_instr} of the following text. 
        Focus on the key points, main ideas, and important conclusions.
        
        TEXT:
        {text[:15000]}  # Limiting text to prevent overloading the model
        
        SUMMARY:
        """
        
        log_activity("AI Request", f"Requesting {length} summary using {model}")
        start_time = time.time()
        
        # Make request to Ollama API
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": model,
                "prompt": prompt,
                "stream": False
            }
        )
        
        end_time = time.time()
        duration = round(end_time - start_time, 2)
        
        if response.status_code == 200:
            summary = response.json()["response"].strip()
            log_activity("AI Response", f"Received summary ({len(summary)} chars) in {duration}s")
            
            # Save result
            if st.session_state.current_file_id:
                if 'summaries' not in st.session_state.analysis_results:
                    st.session_state.analysis_results['summaries'] = {}
                
                st.session_state.analysis_results['summaries'][model] = {
                    'text': summary,
                    'length': length,
                    'timestamp': datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    'duration': duration
                }
            
            return summary
        else:
            error_msg = f"Error from Ollama API: {response.status_code}"
            log_activity("AI Error", error_msg, "ERROR")
            st.error(error_msg)
            return "Error generating summary. Please check if Ollama is running."
    except Exception as e:
        error_msg = f"Error connecting to Ollama: {e}"
        log_activity("AI Connection Error", error_msg, "ERROR")
        st.error(error_msg)
        return "Error connecting to Ollama. Please ensure the service is running."

# Function to create audio from text
def text_to_speech(text):
    try:
        log_activity("TTS", "Generating speech from text")
        start_time = time.time()
        
        tts = gTTS(text=text, lang='en', slow=False)
        audio_fp = BytesIO()
        tts.write_to_fp(audio_fp)
        audio_fp.seek(0)
        
        end_time = time.time()
        duration = round(end_time - start_time, 2)
        log_activity("TTS", f"Generated audio in {duration}s")
        
        return audio_fp
    except Exception as e:
        error_msg = f"Error generating audio: {e}"
        log_activity("TTS Error", error_msg, "ERROR")
        st.error(error_msg)
        return None

# Function to get a more detailed answer from Ollama about a specific query
def get_detailed_answer(text, query, model="llama3"):
    try:
        # Prepare the prompt for answering the query
        prompt = f"""Based on the following document, please answer this question: "{query}"
        
        DOCUMENT:
        {text[:15000]}  # Limiting text to prevent overloading the model
        
        Provide a detailed and accurate answer using information from the document.
        """
        
        log_activity("AI Request", f"Requesting answer to '{query}' using {model}")
        start_time = time.time()
        
        # Make request to Ollama API
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": model,
                "prompt": prompt,
                "stream": False
            }
        )
        
        end_time = time.time()
        duration = round(end_time - start_time, 2)
        
        if response.status_code == 200:
            answer = response.json()["response"].strip()
            log_activity("AI Response", f"Received answer in {duration}s")
            
            # Save to analysis results
            if st.session_state.current_file_id:
                if 'qa' not in st.session_state.analysis_results:
                    st.session_state.analysis_results['qa'] = []
                
                st.session_state.analysis_results['qa'].append({
                    'question': query,
                    'answer': answer,
                    'model': model,
                    'timestamp': datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    'duration': duration
                })
            
            return answer
        else:
            error_msg = f"Error from Ollama API: {response.status_code}"
            log_activity("AI Error", error_msg, "ERROR")
            st.error(error_msg)
            return "Error generating answer. Please check if Ollama is running."
    except Exception as e:
        error_msg = f"Error connecting to Ollama: {e}"
        log_activity("AI Connection Error", error_msg, "ERROR")
        st.error(error_msg)
        return "Error connecting to Ollama. Please ensure the service is running."

# Function to calculate document statistics
def analyze_document_stats(text):
    try:
        # Word count
        words = re.findall(r'\b\w+\b', text)
        word_count = len(words)
        
        # Sentence count
        sentences = re.split(r'[.!?]+', text)
        sentence_count = len([s for s in sentences if s.strip()])
        
        # Average word length
        avg_word_length = sum(len(word) for word in words) / max(1, len(words))
        
        # Average sentence length
        words_per_sentence = word_count / max(1, sentence_count)
        
        # Word frequency
        word_freq = {}
        for word in words:
            word = word.lower()
            if len(word) > 3:  # Only count words longer than 3 chars
                if word not in word_freq:
                    word_freq[word] = 0
                word_freq[word] += 1
        
        # Top 20 most common words
        top_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:20]
        
        # Readability (simple version using average sentence length)
        if words_per_sentence < 12:
            readability = "Easy"
        elif words_per_sentence < 18:
            readability = "Medium"
        else:
            readability = "Complex"
        
        stats = {
            "word_count": word_count,
            "sentence_count": sentence_count,
            "avg_word_length": round(avg_word_length, 2),
            "avg_sentence_length": round(words_per_sentence, 2),
            "top_words": top_words,
            "readability": readability,
            "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        
        log_activity("Analysis", f"Calculated document statistics: {word_count} words, {sentence_count} sentences")
        
        return stats
    except Exception as e:
        error_msg = f"Error analyzing document statistics: {e}"
        log_activity("Analysis Error", error_msg, "ERROR")
        return None

# Function to compare two documents
def compare_documents(text1, text2):
    try:
        # Create TF-IDF vectorizer for comparison
        vectorizer = TfidfVectorizer(stop_words='english')
        tfidf_matrix = vectorizer.fit_transform([text1, text2])
        
        # Calculate similarity
        similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
        
        # Extract unique words
        words1 = set(re.findall(r'\b\w+\b', text1.lower()))
        words2 = set(re.findall(r'\b\w+\b', text2.lower()))
        
        unique_to_1 = words1 - words2
        unique_to_2 = words2 - words1
        common_words = words1.intersection(words2)
        
        # Only keep meaningful words (length > 3)
        unique_to_1 = {w for w in unique_to_1 if len(w) > 3}
        unique_to_2 = {w for w in unique_to_2 if len(w) > 3}
        common_words = {w for w in common_words if len(w) > 3}
        
        # Calculate word count difference
        word_count1 = len(re.findall(r'\b\w+\b', text1))
        word_count2 = len(re.findall(r'\b\w+\b', text2))
        word_diff_pct = abs(word_count1 - word_count2) / max(1, max(word_count1, word_count2)) * 100
        
        comparison = {
            "similarity": round(similarity * 100, 2),
            "word_count1": word_count1,
            "word_count2": word_count2,
            "word_diff_pct": round(word_diff_pct, 2),
            "unique_words1": list(unique_to_1)[:50],  # Limit to top 50
            "unique_words2": list(unique_to_2)[:50],
            "common_words": list(common_words)[:50],
            "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        
        log_activity("Comparison", f"Compared two documents with similarity: {comparison['similarity']}%")
        return comparison
    except Exception as e:
        error_msg = f"Error comparing documents: {e}"
        log_activity("Comparison Error", error_msg, "ERROR")
        return None

# Function to export logs to CSV
def export_logs_to_csv():
    try:
        csv_file = BytesIO()
        fieldnames = ["timestamp", "action", "details", "level"]
        
        writer = csv.DictWriter(csv_file, fieldnames=fieldnames)
        writer.writeheader()
        
        for log in st.session_state.activity_log:
            writer.writerow(log)
        
        csv_file.seek(0)
        log_activity("Export", "Exported activity logs to CSV")
        return csv_file
    except Exception as e:
        error_msg = f"Error exporting logs: {e}"
        log_activity("Export Error", error_msg, "ERROR")
        return None

# Function to export analysis results
def export_analysis_results():
    try:
        zip_buffer = BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED, False) as zip_file:
            # Export summary
            if 'summaries' in st.session_state.analysis_results:
                for model, summary in st.session_state.analysis_results['summaries'].items():
                    summary_txt = summary['text']
                    zip_file.writestr(f"summary_{model}.txt", summary_txt)
            
            # Export Q&A
            if 'qa' in st.session_state.analysis_results:
                qa_data = ""
                for i, qa in enumerate(st.session_state.analysis_results['qa']):
                    qa_data += f"Q{i+1}: {qa['question']}\n\n"
                    qa_data += f"A{i+1}: {qa['answer']}\n\n"
                    qa_data += "---\n\n"
                
                zip_file.writestr("questions_answers.txt", qa_data)
            
            # Export statistics if available
            if 'stats' in st.session_state.analysis_results:
                stats_json = json.dumps(st.session_state.analysis_results['stats'], indent=2)
                zip_file.writestr("document_statistics.json", stats_json)
            
            # Export comparison if available
            if 'comparison' in st.session_state.analysis_results:
                comparison_json = json.dumps(st.session_state.analysis_results['comparison'], indent=2)
                zip_file.writestr("document_comparison.json", comparison_json)
                
            # Export extracted text
            if st.session_state.extracted_text:
                zip_file.writestr("extracted_text.txt", st.session_state.extracted_text)
                
            # Export logs
            logs_csv = export_logs_to_csv()
            if logs_csv:
                zip_file.writestr("activity_logs.csv", logs_csv.getvalue())
        
        zip_buffer.seek(0)
        log_activity("Export", "Exported analysis results to ZIP")
        return zip_buffer
    except Exception as e:
        error_msg = f"Error exporting analysis: {e}"
        log_activity("Export Error", error_msg, "ERROR")
        return None

# Function to generate metrics chart
def generate_metrics_chart(stats):
    if not stats:
        return None
    
    try:
        # Create frequency dataframe
        df = pd.DataFrame(stats['top_words'], columns=['word', 'frequency'])
        
        # Create bar chart
        fig = px.bar(
            df, 
            x='word', 
            y='frequency',
            title='Top Word Frequencies',
            labels={'word': 'Word', 'frequency': 'Frequency'},
            height=400
        )
        
        fig.update_layout(
            xaxis_tickangle=-45,
            margin=dict(l=20, r=20, t=40, b=100)
        )
        
        return fig
    except Exception as e:
        error_msg = f"Error generating chart: {e}"
        log_activity("Chart Error", error_msg, "ERROR")
        return None

# Create a search suggestion function
def get_search_suggestions(text, count=5):
    try:
        # Create a prompt for generating search suggestions
        prompt = f"""Based on the following document, suggest {count} relevant search queries 
        that a user might want to use to find information in this document. 
        Return only the queries as a numbered list, one per line.
        
        DOCUMENT EXTRACT:
        {text[:5000]}
        
        SEARCH SUGGESTIONS:
        """
        
        # Make request to Ollama API
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": st.session_state.settings['default_model'],
                "prompt": prompt,
                "stream": False
            }
        )
        
        if response.status_code == 200:
            suggestions_text = response.json()["response"].strip()
            
            # Extract the suggestions using regex
            suggestions = re.findall(r'\d+\.\s+(.*)', suggestions_text)
            
            # If regex didn't work, just split by lines and clean up
            if not suggestions:
                suggestions = [line.strip() for line in suggestions_text.split('\n') 
                              if line.strip() and not line.strip().startswith('#')]
                
            # Limit to requested count
            suggestions = suggestions[:count]
            
            log_activity("AI", f"Generated {len(suggestions)} search suggestions")
            return suggestions
        else:
            return []
    except Exception as e:
        log_activity("AI Error", f"Error generating search suggestions: {e}", "ERROR")
        return []

# Sidebar for model selection and settings
with st.sidebar:
    st.markdown("### Settings")
    
    # Model selection
    st.session_state.settings['default_model'] = st.selectbox(
        "Default Ollama model:",
        ["llama3", "mistral", "phi", "gemma", "vicuna"],
        index=["llama3", "mistral", "phi", "gemma", "vicuna"].index(st.session_state.settings['default_model'])
    )
    
    # Summary length
    st.session_state.settings['summary_length'] = st.radio(
        "Summary length:",
        ["short", "medium", "long"],
        index=["short", "medium", "long"].index(st.session_state.settings['summary_length'])
    )
    
    # Search settings
    st.session_state.settings['search_results_count'] = st.slider(
        "Max search results:", 
        min_value=3, 
        max_value=15, 
        value=st.session_state.settings['search_results_count']
    )
    
    # UI settings
    st.session_state.settings['highlight_enabled'] = st.checkbox(
        "Highlight search terms", 
        value=st.session_state.settings['highlight_enabled']
    )
    
    st.session_state.settings['save_logs'] = st.checkbox(
        "Save activity logs", 
        value=st.session_state.settings['save_logs']
    )
    
    # Document History
    st.markdown("### Document History")
    
    if st.session_state.document_history:
        for i, doc in enumerate(st.session_state.document_history):
            # Create clickable history items
            if st.button(f"{doc['filename']} ({doc['upload_time']})", key=f"history_{i}"):
                st.session_state.current_file_id = doc['file_id']
                log_activity("History", f"Loaded document from history: {doc['filename']}")
                st.rerun()
    else:
        st.info("No document history yet")
    
    # About section
    st.markdown("### About")
    st.info("""
    Advanced Document Analyzer using Ollama to run open-source LLMs locally.
    Make sure you have Ollama installed and running with the selected model.
    
    To install models, use: `ollama pull modelname`
    """)

# App title and intro
st.markdown("<h1 class='main-header'>📊 Advanced Document Analyzer</h1>", unsafe_allow_html=True)
st.markdown("""
Upload documents to analyze, summarize, search, and compare them using open-source LLMs through Ollama.
This tool provides multiple analysis features and keeps a detailed log of all activities.
""")

# Main interface
uploaded_file = st.file_uploader("Upload a document", type=["pdf", "docx", "txt"])


# Process uploaded file or use file from history
if uploaded_file is not None:
    with st.spinner("Processing document..."):
        # Get file details
        file_details = {"Filename": uploaded_file.name, "Size": uploaded_file.size}
        st.write(f"**File:** {file_details['Filename']} ({file_details['Size']/1000:.1f} KB)")
        
        # Calculate file hash
        file_bytes = uploaded_file.getvalue()
        file_hash = calculate_file_hash(file_bytes)
        
        # Extract text based on file type
        file_extension = Path(uploaded_file.name).suffix.lower()
        
        if file_extension == ".pdf":
            text = extract_text_from_pdf(file_bytes)
        elif file_extension == ".docx":
            text = extract_text_from_docx(file_bytes)
        elif file_extension == ".txt":
            text = extract_text_from_txt(file_bytes)
        else:
            st.error("Unsupported file format")
            text = ""
        
        # Store extracted text and paragraphs in session state
        st.session_state.extracted_text = text
        st.session_state.paragraphs = split_into_paragraphs(text)
        
        # Calculate word count
        word_count = len(text.split())
        
        # Add to document history
        file_id = add_to_history(
            uploaded_file.name,
            file_details["Size"],
            file_hash,
            file_extension,
            word_count
        )
        
        # Set as current file
        st.session_state.current_file_id = file_id
        
        # Reset analysis results for new file
        st.session_state.analysis_results = {}
        
        # Calculate and store document statistics
        stats = analyze_document_stats(text)
        if stats:
            st.session_state.analysis_results['stats'] = stats

# Check if we have a document loaded (either from upload or history)
if st.session_state.extracted_text:
    # Show document length
    word_count = len(st.session_state.extracted_text.split())
    st.write(f"Document contains approximately {word_count} words")
    
    # Create tabs for different features
    tabs = st.tabs([
        "📑 Document Summary", 
        "🔍 Search Document", 
        "📊 Document Analysis", 
        "🔄 Document Comparison",
        "📝 Full Text",
        "📒 Activity Logs"
    ])
    
    # Summary Tab
    with tabs[0]:
        st.markdown("<div class='tab-content'>", unsafe_allow_html=True)
        st.markdown("### Document Summary")
        
        col1, col2 = st.columns([3, 1])
        with col1:
            model = st.selectbox(
                "Select LLM model for summarization:",
                ["llama3", "mistral", "phi", "gemma", "vicuna"],
                index=0,
                key="summary_model"
            )
        
        with col2:
            length = st.selectbox(
                "Summary length:",
                ["short", "medium", "long"],
                index=["short", "medium", "long"].index(st.session_state.settings['summary_length']),
                key="summary_length"
            )
        
        # Button to generate summary
        if st.button("Generate Summary", type="primary", key="btn_summary"):
            with st.spinner(f"Generating summary using {model}..."):
                summary = get_summary_from_ollama(
                    st.session_state.extracted_text, 
                    model=model,
                    length=length
                )
                
                st.markdown("<div class='summary-container'>", unsafe_allow_html=True)
                st.markdown(summary)
                st.markdown("</div>", unsafe_allow_html=True)
                
                # Generate audio summary
                with st.spinner("Generating audio summary..."):
                    audio_bytes = text_to_speech(summary)
                    
                if audio_bytes:
                    st.markdown("### Audio Summary")
                    st.audio(audio_bytes, format="audio/mp3")
                    
                    # Add download button for audio
                    btn = st.download_button(
                        label="Download Audio Summary",
                        data=audio_bytes,
                        file_name=f"summary_{model}.mp3",
                        mime="audio/mp3"
                    )
        
        # Show saved summary if available
        elif 'summaries' in st.session_state.analysis_results and model in st.session_state.analysis_results['summaries']:
            saved_summary = st.session_state.analysis_results['summaries'][model]
            
            st.markdown("<div class='summary-container'>", unsafe_allow_html=True)
            st.markdown(saved_summary['text'])
            st.markdown("</div>", unsafe_allow_html=True)
            
            st.markdown(f"*Generated on {saved_summary['timestamp']} in {saved_summary['duration']}s*")
            
            # Try to find or regenerate audio
            if 'audio' in saved_summary:
                st.markdown("### Audio Summary")
                st.audio(saved_summary['audio'], format="audio/mp3")
            else:
                if st.button("Generate Audio for Summary", key="gen_audio"):
                    with st.spinner("Generating audio summary..."):
                        audio_bytes = text_to_speech(saved_summary['text'])
                        
                    if audio_bytes:
                        st.markdown("### Audio Summary")
                        st.audio(audio_bytes, format="audio/mp3")
                        
                        # Add download button for audio
                        btn = st.download_button(
                            label="Download Audio Summary",
                            data=audio_bytes,
                            file_name=f"summary_{model}.mp3",
                            mime="audio/mp3"
                        )
                        
                        # Save audio to avoid regenerating
                        st.session_state.analysis_results['summaries'][model]['audio'] = audio_bytes
        else:
            st.info("Click 'Generate Summary' to create a summary of the document using the selected model.")
        
        st.markdown("</div>", unsafe_allow_html=True)
    
    # Search Tab
    with tabs[1]:
        st.markdown("<div class='tab-content'>", unsafe_allow_html=True)
        st.markdown("### Search Document Content")
        
        # Get AI-generated search suggestions
        if st.session_state.paragraphs and len(st.session_state.paragraphs) > 0:
            if 'search_suggestions' not in st.session_state.analysis_results:
                if st.button("Generate Search Suggestions", key="gen_suggestions"):
                    with st.spinner("Generating search suggestions..."):
                        suggestions = get_search_suggestions(st.session_state.extracted_text, count=5)
                        if suggestions:
                            st.session_state.analysis_results['search_suggestions'] = suggestions
                            st.rerun()
            
            # Display suggestions if available
            if 'search_suggestions' in st.session_state.analysis_results:
                st.markdown("**Try these search queries:**")
                for i, suggestion in enumerate(st.session_state.analysis_results['search_suggestions']):
                    if st.button(suggestion, key=f"suggestion_{i}"):
                        search_query = suggestion
                        st.session_state.search_query = suggestion
                        st.rerun()
        
        # Search input
        if 'search_query' not in st.session_state:
            st.session_state.search_query = ""
            
        search_query = st.text_input(
            "Enter search terms or questions:", 
            value=st.session_state.search_query,
            placeholder="Enter keywords or a question about the document..."
        )
        
        search_col1, search_col2 = st.columns([1, 3])
        
        with search_col1:
            search_type = st.radio("Search Type:", ["Keyword Search", "Ask Question"])
            max_results = st.slider(
                "Results to show:",
                min_value=3,
                max_value=15,
                value=st.session_state.settings['search_results_count']
            )
        
        with search_col2:
            if st.button("Search", type="primary", key="btn_search"):
                if search_query:
                    if search_type == "Keyword Search":
                        # Perform keyword search
                        with st.spinner("Searching document..."):
                            results = search_document(search_query, st.session_state.paragraphs, top_k=max_results)
                            st.session_state.search_results = results
                            
                            # Store search in history
                            if 'search_history' not in st.session_state.analysis_results:
                                st.session_state.analysis_results['search_history'] = []
                            
                            st.session_state.analysis_results['search_history'].append({
                                'type': 'keyword',
                                'query': search_query,
                                'timestamp': datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                'results_count': len(results)
                            })
                    else:  # Ask Question
                        # Get detailed answer from LLM
                        with st.spinner("Generating answer..."):
                            answer = get_detailed_answer(
                                st.session_state.extracted_text, 
                                search_query, 
                                model=st.session_state.settings['default_model']
                            )
                            st.markdown("### Answer")
                            st.markdown(f"<div class='summary-container'>{answer}</div>", unsafe_allow_html=True)
                            
                            # Also perform keyword search to show relevant sections
                            results = search_document(search_query, st.session_state.paragraphs)
                            st.session_state.search_results = results
        
        # Display search results
        if st.session_state.search_results:
            st.markdown("### Search Results")
            for i, result in enumerate(st.session_state.search_results):
                highlighted_text = highlight_text(result["text"], search_query)
                st.markdown(f"""
                <div class='search-result'>
                    <strong>Match {i+1}</strong> (Relevance: {result['score']:.2f})<br>
                    {highlighted_text}
                </div>
                """, unsafe_allow_html=True)
                
            if not st.session_state.search_results:
                st.info("No matching content found. Try different search terms.")
        
        st.markdown("</div>", unsafe_allow_html=True)
    
    # Analysis Tab
    with tabs[2]:
        st.markdown("<div class='tab-content'>", unsafe_allow_html=True)
        st.markdown("### Document Analysis")
        
        if 'stats' in st.session_state.analysis_results:
            stats = st.session_state.analysis_results['stats']
            
            # Display metrics in cards
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.markdown("""
                <div class='metrics-card'>
                    <div class='metrics-label'>Word Count</div>
                    <div class='metrics-value'>{}</div>
                </div>
                """.format(stats['word_count']), unsafe_allow_html=True)
            
            with col2:
                st.markdown("""
                <div class='metrics-card'>
                    <div class='metrics-label'>Sentences</div>
                    <div class='metrics-value'>{}</div>
                </div>
                """.format(stats['sentence_count']), unsafe_allow_html=True)
            
            with col3:
                st.markdown("""
                <div class='metrics-card'>
                    <div class='metrics-label'>Avg Word Length</div>
                    <div class='metrics-value'>{}</div>
                </div>
                """.format(stats['avg_word_length']), unsafe_allow_html=True)
            
            with col4:
                st.markdown("""
                <div class='metrics-card'>
                    <div class='metrics-label'>Readability</div>
                    <div class='metrics-value'>{}</div>
                </div>
                """.format(stats['readability']), unsafe_allow_html=True)
            
            # Display word frequency chart
            st.markdown("### Word Frequency Analysis")
            chart = generate_metrics_chart(stats)
            if chart:
                st.plotly_chart(chart, use_container_width=True)
            
            # Topic extraction
            if st.button("Extract Key Topics", key="extract_topics"):
                with st.spinner("Analyzing document topics..."):
                    prompt = f"""Analyze the following document and extract the 5-7 main topics or themes. 
                    For each topic, provide a concise label and a brief description.
                    
                    DOCUMENT:
                    {st.session_state.extracted_text[:10000]}
                    
                    FORMAT YOUR RESPONSE AS:
                    Topic 1: [Label]
                    [Brief description]
                    
                    Topic 2: [Label]
                    [Brief description]
                    
                    And so on.
                    """
                    
                    # Make request to Ollama API
                    response = requests.post(
                        "http://localhost:11434/api/generate",
                        json={
                            "model": st.session_state.settings['default_model'],
                            "prompt": prompt,
                            "stream": False
                        }
                    )
                    
                    if response.status_code == 200:
                        topics = response.json()["response"].strip()
                        
                        st.markdown("### Key Topics")
                        st.markdown(f"<div class='summary-container'>{topics}</div>", unsafe_allow_html=True)
                        
                        # Save to analysis results
                        st.session_state.analysis_results['topics'] = {
                            'text': topics,
                            'timestamp': datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        }
            
            # Display saved topics if available
            elif 'topics' in st.session_state.analysis_results:
                st.markdown("### Key Topics")
                st.markdown(f"<div class='summary-container'>{st.session_state.analysis_results['topics']['text']}</div>", unsafe_allow_html=True)
        else:
            st.info("Document statistics are not available. Try uploading a new document.")
        
        st.markdown("</div>", unsafe_allow_html=True)
    
    # Comparison Tab
    with tabs[3]:
        st.markdown("<div class='tab-content'>", unsafe_allow_html=True)
        st.markdown("### Document Comparison")
        
        st.info("Upload another document to compare with the current one.")
        
        comparison_file = st.file_uploader("Upload document for comparison", type=["pdf", "docx", "txt"], key="comparison_upload")
        
        if comparison_file:
            with st.spinner("Processing comparison document..."):
                # Extract text from comparison file
                comp_file_bytes = comparison_file.getvalue()
                comp_file_extension = Path(comparison_file.name).suffix.lower()
                
                if comp_file_extension == ".pdf":
                    comparison_text = extract_text_from_pdf(comp_file_bytes)
                elif comp_file_extension == ".docx":
                    comparison_text = extract_text_from_docx(comp_file_bytes)
                elif comp_file_extension == ".txt":
                    comparison_text = extract_text_from_txt(comp_file_bytes)
                else:
                    st.error("Unsupported file format")
                    comparison_text = ""
                
                if comparison_text:
                    # Perform comparison
                    comparison = compare_documents(st.session_state.extracted_text, comparison_text)
                    
                    if comparison:
                        # Store comparison results
                        st.session_state.analysis_results['comparison'] = comparison
                        
                        # Display similarity
                        st.markdown("### Similarity Analysis")
                        
                        # Create similarity gauge
                        fig = go.Figure(go.Indicator(
                            mode = "gauge+number",
                            value = comparison['similarity'],
                            title = {'text': "Document Similarity"},
                            gauge = {
                                'axis': {'range': [0, 100]},
                                'bar': {'color': "darkblue"},
                                'steps': [
                                    {'range': [0, 30], 'color': "lightgray"},
                                    {'range': [30, 70], 'color': "gray"},
                                    {'range': [70, 100], 'color': "darkgray"}
                                ],
                                'threshold': {
                                    'line': {'color': "red", 'width': 4},
                                    'thickness': 0.75,
                                    'value': 90
                                }
                            }
                        ))
                        
                        fig.update_layout(height=250)
                        st.plotly_chart(fig, use_container_width=True)
                        
                        # Display metrics
                        col1, col2, col3 = st.columns(3)
                        
                        with col1:
                            st.metric("Original Document", f"{comparison['word_count1']} words")
                        
                        with col2:
                            st.metric("Comparison Document", f"{comparison['word_count2']} words")
                        
                        with col3:
                            st.metric("Word Count Difference", f"{comparison['word_diff_pct']}%")
                        
                        # Word comparison
                        word_tabs = st.tabs(["Unique to Original", "Unique to Comparison", "Common Words"])
                        
                        with word_tabs[0]:
                            st.markdown("Words that appear only in the original document:")
                            st.write(", ".join(comparison['unique_words1']))
                        
                        with word_tabs[1]:
                            st.markdown("Words that appear only in the comparison document:")
                            st.write(", ".join(comparison['unique_words2']))
                        
                        with word_tabs[2]:
                            st.markdown("Common words between both documents:")
                            st.write(", ".join(comparison['common_words']))
                            
                        # Detailed analysis button
                        if st.button("Generate Detailed Comparison", key="detailed_comparison"):
                            with st.spinner("Analyzing documents..."):
                                prompt = f"""Compare these two documents and provide a detailed analysis of their similarities and differences.
                                Focus on content, style, tone, and key points. Identify significant variations.
                                
                                DOCUMENT 1:
                                {st.session_state.extracted_text[:7500]}
                                
                                DOCUMENT 2:
                                {comparison_text[:7500]}
                                
                                ANALYSIS:
                                """
                                
                                # Make request to Ollama API
                                response = requests.post(
                                    "http://localhost:11434/api/generate",
                                    json={
                                        "model": st.session_state.settings['default_model'],
                                        "prompt": prompt,
                                        "stream": False
                                    }
                                )
                                
                                if response.status_code == 200:
                                    analysis = response.json()["response"].strip()
                                    
                                    st.markdown("### Detailed Analysis")
                                    st.markdown(f"<div class='summary-container'>{analysis}</div>", unsafe_allow_html=True)
                                    
                                    # Save to comparison results
                                    st.session_state.analysis_results['comparison']['detailed_analysis'] = {
                                        'text': analysis,
                                        'timestamp': datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                                    }
                        
                        # Show saved detailed analysis if available
                        elif 'detailed_analysis' in st.session_state.analysis_results.get('comparison', {}):
                            st.markdown("### Detailed Analysis")
                            st.markdown(
                                f"<div class='summary-container'>{st.session_state.analysis_results['comparison']['detailed_analysis']['text']}</div>", 
                                unsafe_allow_html=True
                            )
        
        st.markdown("</div>", unsafe_allow_html=True)
    
    # Full Text Tab
    with tabs[4]:
        st.markdown("<div class='tab-content'>", unsafe_allow_html=True)
        st.markdown("### Document Text")
        st.text_area("", st.session_state.extracted_text, height=500)
        st.markdown("</div>", unsafe_allow_html=True)
    
    # Activity Logs Tab
    with tabs[5]:
        st.markdown("<div class='tab-content'>", unsafe_allow_html=True)
        st.markdown("### Activity Logs")
        
        # Actions for logs
        log_col1, log_col2 = st.columns([1, 3])
        
        with log_col1:
            if st.button("Clear Logs", key="clear_logs"):
                st.session_state.activity_log = []
                log_activity("Logs", "Cleared activity logs")
                st.rerun()
        
        with log_col2:
            csv_logs = export_logs_to_csv()
            if csv_logs:
                st.download_button(
                    label="Download Logs as CSV",
                    data=csv_logs,
                    file_name="document_analyzer_logs.csv",
                    mime="text/csv",
                    key="download_logs"
                )
        
        # Display log entries
        st.markdown("#### Log Entries")
        for log in reversed(st.session_state.activity_log):
            log_class = f"log-{log['level'].lower()}" if log['level'] in ["INFO", "WARNING", "ERROR"] else "log-info"
            st.markdown(f"""
            <div class='log-entry'>
                <span class='{log_class}'>[{log['level']}]</span> <strong>{log['timestamp']}</strong> - {log['action']}: {log['details']}
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown("</div>", unsafe_allow_html=True)
    
    # Export options
    st.markdown("---")
    st.markdown("### Export Results")
    
    export_col1, export_col2 = st.columns(2)
    
    with export_col1:
        if st.session_state.extracted_text:
            zip_buffer = export_analysis_results()
            if zip_buffer:
                st.download_button(
                    label="Download All Analysis Results",
                    data=zip_buffer,
                    file_name=f"document_analysis_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                    mime="application/zip",
                    key="download_all"
                )
    
    with export_col2:
        if 'summaries' in st.session_state.analysis_results:
            for model, summary in st.session_state.analysis_results['summaries'].items():
                st.download_button(
                    label=f"Download {model.capitalize()} Summary (Text)",
                    data=summary['text'],
                    file_name=f"summary_{model}.txt",
                    mime="text/plain",
                    key=f"download_summary_{model}"
                )
else:
    st.info("Please upload a document to get started or select one from the history.")

# Instructions for Ollama
st.markdown("---")
with st.expander("How to set up Ollama"):
    st.markdown("""
    ### Setting up Ollama
    
    1. Install Ollama from [ollama.ai](https://ollama.ai)
    2. Pull the model you want to use:
       ```
       ollama pull llama3
       ```
    3. Ensure Ollama is running in the background
    4. This app connects to Ollama's API at http://localhost:11434
    
    If you encounter issues, make sure Ollama is properly installed and running.
    """)

# Footer
st.markdown("---")
st.markdown("Advanced Document Analyzer | Built with Streamlit, Ollama, and open-source LLMs")
